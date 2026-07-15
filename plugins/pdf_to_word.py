import os
import fitz
import tempfile
import logging

from PIL import Image

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger(__name__)

OUTPUT_FOLDER = "outputs"

TEMP_FOLDER = tempfile.gettempdir()

os.makedirs(
    OUTPUT_FOLDER,
    exist_ok=True
)


def create_document(title):

    doc = Document()

    heading = doc.add_heading(
        title,
        level=1
    )

    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return doc


def add_text(doc, text):

    if not text.strip():

        return

    paragraphs = text.split("\n")

    for p in paragraphs:

        p = p.strip()

        if p:

            doc.add_paragraph(p)


def extract_images(page):

    images = []

    for img in page.get_images(full=True):

        xref = img[0]

        images.append(xref)

    return images

def save_image(pdf, xref, page_no, index):

    image = pdf.extract_image(xref)

    image_bytes = image["image"]

    extension = image["ext"]

    filename = os.path.join(

        TEMP_FOLDER,

        f"page_{page_no}_{index}.{extension}"

    )

    with open(filename, "wb") as f:

        f.write(image_bytes)

    return filename


def add_images(doc, pdf, page):

    image_list = extract_images(page)

    if len(image_list) == 0:

        return

    doc.add_paragraph("")

    for i, xref in enumerate(image_list):

        try:

            image_path = save_image(

                pdf,

                xref,

                page.number,

                i

            )

            try:

                img = Image.open(image_path)

                width, height = img.size

                img.close()

                if width > height:

                    doc.add_picture(

                        image_path,

                        width=Inches(6)

                    )

                else:

                    doc.add_picture(

                        image_path,

                        width=Inches(3)

                    )

            except:

                doc.add_picture(

                    image_path,

                    width=Inches(5)

                )

            os.remove(image_path)

            doc.add_paragraph("")

        except Exception as ex:

            logger.warning(

                f"Image skipped: {ex}"

            )

def process_page(doc, pdf, page):

    logger.info(

        f"Processing page {page.number + 1}"

    )

    text = page.get_text("text")

    if text.strip():

        add_text(

            doc,

            text

        )

    add_images(

        doc,

        pdf,

        page
    )

    doc.add_page_break()


def convert_pdf_to_word(

    pdf_path,

    original_name

):

    output_path = os.path.join(

        OUTPUT_FOLDER,

        original_name + ".docx"

    )

    pdf = fitz.open(pdf_path)

    document = create_document(

        original_name

    )

    try:

        total_pages = len(pdf)

        logger.info(

            f"Pages: {total_pages}"

        )

        for page in pdf:

            process_page(

                document,

                pdf,

                page

            )

        document.save(

            output_path

        )

        logger.info(

            "Word document created successfully."

        )

        return output_path

    except Exception as ex:

        logger.exception(ex)

        raise Exception(

            "Unable to convert PDF to Word."

        )

    finally:

        pdf.close()